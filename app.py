import streamlit as st
import os
import tempfile
import pbixray
from docx import Document

# Define the functions within the Streamlit script for cohesion

def extract_pbix_info(pbix_file_path):
    """
    Extracts detailed information from a PBIX file.

    Args:
        pbix_file_path: Path to the PBIX file.

    Returns:
        A dictionary containing extracted information, or None if an error occurs.
    """
    try:
        # Ensure the file exists
        if not os.path.exists(pbix_file_path):
            st.error(f"Error: File not found at {pbix_file_path}")
            return None

        with pbixray.open(pbix_file_path) as pbix:
            info = {}

            # Extract data model information
            info["data_model"] = {
                "tables": [table.Name for table in pbix.model.Tables],
                "relationships": [rel.Name for rel in pbix.model.Relationships],
                "measures": [measure.Name for measure in pbix.model.Measures]
            }

            # Extract page information
            info["pages"] = [page.DisplayName for page in pbix.report.Pages]

            # Extract visuals information (simplified example, can be expanded)
            visuals_info = {}
            if pbix.report.Pages:
                first_page = pbix.report.Pages[0]
                visuals_info[first_page.DisplayName] = "Details about visuals would go here (requires deeper parsing)"

            info["visuals"] = visuals_info

            return info

    except Exception as e:
        st.error(f"An error occurred during PBIX analysis: {e}")
        return None

def create_word_documentation(pbix_info, output_path):
    """
    Generates a Word document from extracted PBIX information.

    Args:
        pbix_info: Dictionary containing extracted PBIX information.
        output_path: Path to save the generated Word document.
    """
    document = Document()

    # Add title
    document.add_heading('PBIX File Documentation', 0)

    # Iterate through extracted information
    for section, details in pbix_info.items():
        # Add section heading
        document.add_heading(section.replace('_', ' ').title(), level=1)

        if isinstance(details, list):
            # If details are a list, add them as list items
            for item in details:
                document.add_paragraph(str(item), style='List Bullet')
        elif isinstance(details, dict):
            # If details are a dictionary, add sub-headings and paragraphs
            for sub_section, sub_details in details.items():
                document.add_heading(sub_section.replace('_', ' ').title(), level=2)
                if isinstance(sub_details, list):
                    for item in sub_details:
                        document.add_paragraph(str(item), style='List Bullet')
                else:
                    document.add_paragraph(str(sub_details))
        else:
            # For other types of details, add as a paragraph
            document.add_paragraph(str(details))

    # Add a note about potentially missing detailed information
    document.add_paragraph(
        "Note: Detailed information for some sections (e.g., specific visual configurations)"
        " may require deeper parsing and are represented with placeholders if not fully extracted."
    )

    # Save the document
    document.save(output_path)
    st.success(f"Word document saved to {output_path}")


# Streamlit App Layout and Logic
st.set_page_config(layout="wide")
st.title("PBIX File Analyzer and Document Generator")

uploaded_file = st.file_uploader("Upload a PBIX file", type="pbix")

if uploaded_file is not None:
    # Create a temporary file to save the uploaded PBIX
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pbix") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_file_path = tmp_file.name

    st.subheader("Analyzing PBIX file...")

    # Extract information
    pbix_info = extract_pbix_info(tmp_file_path)

    if pbix_info:
        st.subheader("Extracted Information:")

        # Display extracted information
        st.write("### Data Model")
        st.json(pbix_info.get("data_model", {}))

        st.write("### Pages")
        st.json(pbix_info.get("pages", []))

        st.write("### Visuals (Simplified)")
        st.json(pbix_info.get("visuals", {}))

        # Add a button to generate documentation
        if st.button("Generate Word Documentation"):
            doc_output_path = tmp_file_path.replace(".pbix", "_documentation.docx")
            try:
                create_word_documentation(pbix_info, doc_output_path)

                # Provide download button for the generated document
                if os.path.exists(doc_output_path):
                    with open(doc_output_path, "rb") as doc_file:
                        st.download_button(
                            label="Download Word Document",
                            data=doc_file,
                            file_name=os.path.basename(doc_output_path),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.error("Documentation file was not found after generation.")

            except Exception as e:
                st.error(f"Error generating documentation: {e}")

    else:
        st.error("Could not extract information from the PBIX file. Please check the file and try again.")

    # Clean up the temporary file - placed outside the if pbix_info block
    # to ensure cleanup even if extraction fails
    try:
        if os.path.exists(tmp_file_path):
            os.unlink(tmp_file_path)
            # st.info(f"Cleaned up temporary file: {tmp_file_path}") # Optional: for debugging
    except Exception as e:
        st.warning(f"Could not clean up temporary file {tmp_file_path}: {e}")

else:
    st.info("Please upload a PBIX file to get started.")
