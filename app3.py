import streamlit as st
import pbixray
import os
import tempfile
import io
import pandas as pd
import traceback
import xlsxwriter
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def handle_file_upload():
    """Handles file upload and returns the file object."""
    st.sidebar.header("Upload PBIX File")
    uploaded_file = st.sidebar.file_uploader("Choose a .pbix file", type="pbix")
    return uploaded_file

def process_pbix_file(uploaded_file):
    """Processes the PBIX file using pbixray and extracts data."""
    if uploaded_file is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pbix") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name

        try:
            st.write("Processing PBIX file...")
            unpacker = pbixray.pbix_unpacker.PbixUnpacker(tmp_file_path)

            extracted_data = {}
            # Removed lines that access attributes not present on PbixUnpacker
            # extracted_data["Metadata"] = unpacker.metadata
            # extracted_data["Schema"] = unpacker.schema
            # extracted_data["Relationships"] = unpacker.relationships
            # extracted_data["Power Query"] = unpacker.power_query
            # extracted_data["M Parameters"] = unpacker.m_parameters
            # extracted_data["DAX Tables"] = unpacker.dax_tables
            # extracted_data["DAX Measures"] = unpacker.dax_measures

            # Access the data_model with error handling
            try:
                extracted_data["Data Model"] = unpacker.data_model
                st.write("Successfully accessed Data Model object (might not contain full model structure).")
                # Add some debugging info about the data_model structure
                if extracted_data["Data Model"]:
                    st.write(f"Type of unpacker.data_model: {type(extracted_data['Data Model'])}")
                    # Introspect the data_model object to find the tables attribute
                    st.write("Introspecting Data Model object attributes:")
                    st.write([attr for attr in dir(extracted_data["Data Model"]) if not attr.startswith('_')])

            except AttributeError as ae:
                st.error(f"AttributeError accessing data_model or its attributes: {ae}")
                extracted_data["Data Model"] = None # Ensure Data Model is None on error
            except Exception as e:
                st.error(f"An unexpected error occurred accessing data_model: {e}")
                extracted_data["Data Model"] = None # Ensure Data Model is None on error


            # Extract actual data for tables
            table_data = {}
            # Only attempt to process tables if data_model and its components are available
            # Corrected access to tables directly from data_model, based on previous introspection
            if extracted_data["Data Model"] and hasattr(extracted_data["Data Model"], 'tables') and extracted_data["Data Model"].tables:
                 for table in extracted_data["Data Model"].tables:
                    try:
                        # As determined in the previous step, pbixray.data_model does not support
                        # direct extraction of table data into DataFrames.
                        # We will store a placeholder indicating this limitation.
                        table_data[table.name] = pd.DataFrame({"Status": [f"Data extraction for {table.name} not supported by pbixray.data_model for direct viewing."]})

                    except Exception as data_e:
                        table_data[table.name] = pd.DataFrame({"Error": [f"Could not extract data for {table.name}: {data_e}"]})

            extracted_data["Table Data"] = table_data # Store the extracted (or placeholder) data

            st.success("PBIX file processed successfully!")
            return extracted_data

        except Exception as e:
            st.error(f"Error processing PBIX file: {e}")
            st.error(traceback.format_exc())
            return None
        finally:
            # Ensure temporary file is removed even if an error occurs
            if 'tmp_file_path' in locals() and os.path.exists(tmp_file_path):
                os.remove(tmp_file_path)
    return None

def display_extracted_data(extracted_data):
    """Displays the extracted data in a generic way."""
    if extracted_data:
        st.header("Extracted Information")
        for section, data in extracted_data.items():
            # Skip Data Model and Table Data here, as they are handled separately
            if section in ["Data Model", "Table Data"]:
                continue
            st.subheader(section)
            if isinstance(data, (dict, list)):
                st.json(data)
            else:
                st.write(data)

# The generate_excel_doc function from the original code, kept for potential future use
def generate_excel_doc(data):
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    worksheet = workbook.add_worksheet("PBIX Data")

    row = 0
    for section, content in data.items():
        worksheet.write(row, 0, section)
        row += 1
        if isinstance(content, (dict, list)):
            try:
                # Attempt to flatten complex structures for Excel
                df = pd.json_normalize(content)
                df.to_excel(workbook, sheet_name="PBIX Data", startrow=row, index=False, header=True)
                row += len(df) + 2 # Add some space after the table
            except Exception as e:
                worksheet.write(row, 1, f"Could not display complex data: {e}")
                row += 2
        # Handle DataFrame specifically if we were able to extract table data
        elif isinstance(content, pd.DataFrame):
             content.to_excel(workbook, sheet_name="PBIX Data", startrow=row, index=False, header=True)
             row += len(content) + 2
        else:
            worksheet.write(row, 1, str(content))
            row += 2

    workbook.close()
    output.seek(0)
    return output

def generate_word_doc(data):
    """Generates a Word document from extracted data."""
    document = Document()
    document.add_heading('PBIX Extracted Data', 0)

    for section, content in data.items():
        document.add_heading(section, level=1)
        if isinstance(content, (dict, list)):
            # Convert complex data to a string representation
            document.add_paragraph(str(content))
        elif isinstance(content, pd.DataFrame):
            # Attempt to add DataFrame to Word
            document.add_paragraph(content.to_string()) # Convert DataFrame to string for simplicity
        else:
            document.add_paragraph(str(content))

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output

def generate_pdf_doc(data):
    """Generates a PDF document from extracted data."""
    output = io.BytesIO()
    c = canvas.Canvas(output, pagesize=letter)
    width, height = letter

    c.drawString(100, height - 50, "PBIX Extracted Data")
    y_position = height - 100

    for section, content in data.items():
        c.drawString(100, y_position, f"Section: {section}")
        y_position -= 20
        if isinstance(content, (dict, list)):
            # Convert complex data to a string representation for PDF
            content_str = str(content)
            # Split content string if it's too long to fit on one line
            lines = []
            current_line = ""
            for word in content_str.split():
                if c.stringWidth(current_line + word, 'Helvetica', 12) < width - 200:
                    current_line += word + " "
                else:
                    lines.append(current_line)
                    current_line = word + " "
            lines.append(current_line)

            for line in lines:
                c.drawString(120, y_position, line)
                y_position -= 15
                if y_position < 50: # Check if new page is needed
                    c.showPage()
                    c.drawString(100, height - 50, "PBIX Extracted Data (cont.)")
                    y_position = height - 100
        elif isinstance(content, pd.DataFrame):
             # Convert DataFrame to string and add to PDF
             content_str = content.to_string()
             lines = content_str.split('\n')
             for line in lines:
                 c.drawString(120, y_position, line)
                 y_position -= 15
                 if y_position < 50: # Check if new page is needed
                     c.showPage()
                     c.drawString(100, height - 50, "PBIX Extracted Data (cont.)")
                     y_position = height - 100
        else:
            c.drawString(120, y_position, str(content))
            y_position -= 15

        y_position -= 20 # Space between sections
        if y_position < 50: # Check if new page is needed before next section
            c.showPage()
            c.drawString(100, height - 50, "PBIX Extracted Data (cont.)")
            y_position = height - 100


    c.save()
    output.seek(0)
    return output


def main():
    st.set_page_config(page_title="PBIX Ray Explorer", layout="wide")

    st.title("PBIX Ray Explorer")

    uploaded_file = handle_file_upload()

    if uploaded_file:
        extracted_data = process_pbix_file(uploaded_file)

        if extracted_data:
            # Display schema details from Data Model
            st.header("Schema Details (from Data Model)")
            data_model = extracted_data.get("Data Model")

            # Corrected access to tables directly from data_model, based on previous introspection
            if data_model and hasattr(data_model, 'tables') and data_model.tables:
                tables = data_model.tables
                schema_details = {}
                schema_names = []
                for table in tables:
                    # Assuming table objects have 'name' and 'columns' attributes
                    if hasattr(table, 'name') and hasattr(table, 'columns'):
                        schema_details[table.name] = {
                            "Columns": [{
                                "Name": col.name,
                                "DataType": col.data_type,
                                "isHidden": col.is_hidden,
                                "FormatString": col.format_string
                            } for col in table.columns] if hasattr(table, 'columns') else [],
                            "Measures": [{
                                "Name": measure.name,
                                "Expression": measure.expression,
                                "isHidden": measure.is_hidden,
                                "FormatString": measure.format_string
                            } for measure in table.measures] if hasattr(table, 'measures') else []
                        }
                        schema_names.append(table.name)
                    else:
                        st.warning(f"Skipping unexpected table object structure: {type(table)}")


                if schema_names:
                    st.json(schema_details)

                    # Add selectbox for schema selection
                    st.subheader("Select a Schema to View Data")
                    selected_schema = st.selectbox("Choose a schema", ["--Select--"] + schema_names)

                    # Display selected schema data (Still a placeholder due to pbixray limitations)
                    if selected_schema and selected_schema != "--Select--":
                         st.subheader(f"Data for Schema: {selected_schema}")
                         st.info(f"Data extraction for schema '{selected_schema}' is not directly supported by the current pbixray data model object for viewing.")

            else:
                st.info("Could not extract Data Model details or no tables were found. Ensure the uploaded PBIX file is valid and contains a data model.")


            # Display other extracted data (excluding Data Model)
            # Keep Table Data in extracted_data but don't display here as it's placeholder
            display_extracted_data({k: v for k, v in extracted_data.items() if k not in ["Data Model", "Table Data"]})

            # Add download buttons
            st.sidebar.header("Download Output")
            if extracted_data:
                # Pass the entire extracted_data to the generators, they can decide what to include
                excel_output = generate_excel_doc(extracted_data)
                st.sidebar.download_button(label="Download Data as Excel",
                                        data=excel_output,
                                        file_name="pbix_data.xlsx",
                                        mime="application/vnd.openxmlformats-officedocument.spreadsheet.sheet")

                word_output = generate_word_doc(extracted_data)
                st.sidebar.download_button(label="Download Data as Word",
                                        data=word_output,
                                        file_name="pbix_data.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                pdf_output = generate_pdf_doc(extracted_data)
                st.sidebar.download_button(label="Download Data as PDF",
                                        data=pdf_output,
                                        file_name="pbix_data.pdf",
                                        mime="application/pdf")


if __name__ == "__main__":
    main()
