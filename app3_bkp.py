import streamlit as st
import pbixray
import os
import tempfile
import io
import pandas as pd
import traceback
import xlsxwriter
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def handle_file_upload():
    """Handles file upload and returns the file object."""
    st.sidebar.header("Upload PBIX File")
    uploaded_file = st.sidebar.file_uploader("Choose a .pbix file", type="pbix")
    return uploaded_file

def process_pbix_file(uploaded_file):
    """Processes the PBIX file using pbixray and extracts data."""
    if uploaded_file is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pbix") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name

        try:
            st.write("Processing PBIX file...")
            unpacker = pbixray.pbix_unpacker.PbixUnpacker(tmp_file_path)

            extracted_data = {}
            # Access attributes directly from the PbixUnpacker object
            # Add error handling for each attribute access as they might not always be present
            try:
                extracted_data["Metadata"] = unpacker.metadata
            except AttributeError:
                extracted_data["Metadata"] = "Not available"
            try:
                extracted_data["Schema"] = unpacker.schema
            except AttributeError:
                extracted_data["Schema"] = "Not available"
            try:
                extracted_data["Relationships"] = unpacker.relationships
            except AttributeError:
                extracted_data["Relationships"] = "Not available"
            try:
                extracted_data["Power Query"] = unpacker.power_query
            except AttributeError:
                extracted_data["Power Query"] = "Not available"
            try:
                extracted_data["M Parameters"] = unpacker.m_parameters
            except AttributeError:
                extracted_data["M Parameters"] = "Not available"
            try:
                extracted_data["DAX Tables"] = unpacker.dax_tables
            except AttributeError:
                extracted_data["DAX Tables"] = "Not available"
            try:
                extracted_data["DAX Measures"] = unpacker.dax_measures
            except AttributeError:
                extracted_data["DAX Measures"] = "Not available"

            # The 'data_model' attribute provides lower-level access and doesn't have
            # the structured schema information directly. We will not use it for schema display.
            # extracted_data["Data Model"] = unpacker.data_model # Not using for schema/table listing

            st.success("PBIX file processed successfully!")
            return extracted_data

        except Exception as e:
            st.error(f"Error processing PBIX file: {e}")
            st.error(traceback.format_exc())
            return None
        finally:
            # Ensure temporary file is removed even if an error occurs
            if 'tmp_file_path' in locals() and os.path.exists(tmp_file_path):
                os.remove(tmp_file_path)
    return None

def display_extracted_data(extracted_data):
    """Displays the extracted data in a generic way."""
    if extracted_data:
        st.header("Extracted Information")
        for section, data in extracted_data.items():
            st.subheader(section)
            if isinstance(data, (dict, list)):
                st.json(data)
            else:
                st.write(data)

# The generate_excel_doc function from the original code, kept for potential future use
def generate_excel_doc(data):
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    # Use a single worksheet for simplicity or iterate through data sections
    worksheet = workbook.add_worksheet("PBIX Data Summary")

    row = 0
    for section, content in data.items():
        worksheet.write(row, 0, section)
        row += 1
        if isinstance(content, (dict, list)):
            try:
                # Attempt to flatten complex structures for Excel
                df = pd.json_normalize(content)
                # Write to a new sheet for each section if feasible, or flatten
                # For simplicity, writing string representation to the summary sheet
                worksheet.write(row, 1, str(content))
                row += 1
                # If writing to separate sheets:
                # sheet_name = section[:31].replace(' ', '_').replace('-', '_') # Sanitize sheet name
                # df.to_excel(workbook, sheet_name=sheet_name, index=False)
            except Exception as e:
                worksheet.write(row, 1, f"Could not display complex data: {e}")
                row += 1
        # Handle DataFrame specifically if we were able to extract table data (not currently supported by pbixray.data_model)
        elif isinstance(content, pd.DataFrame):
             # Convert DataFrame to string for simplicity in summary sheet
             worksheet.write(row, 1, content.to_string())
             row += 1
        else:
            worksheet.write(row, 1, str(content))
            row += 1
        row += 1 # Add a blank row between sections

    workbook.close()
    output.seek(0)
    return output

def generate_word_doc(data):
    """Generates a Word document from extracted data."""
    document = Document()
    document.add_heading('PBIX Extracted Data', 0)

    for section, content in data.items():
        document.add_heading(section, level=1)
        if isinstance(content, (dict, list)):
            # Convert complex data to a string representation
            document.add_paragraph(str(content))
        elif isinstance(content, pd.DataFrame):
            # Attempt to add DataFrame to Word (as string)
            document.add_paragraph(content.to_string())
        else:
            document.add_paragraph(str(content))

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output

def generate_pdf_doc(data):
    """Generates a PDF document from extracted data."""
    output = io.BytesIO()
    c = canvas.Canvas(output, pagesize=letter)
    width, height = letter
    margin = 50
    y_position = height - margin
    line_height = 14
    section_spacing = 20

    def draw_string_with_wrap(text, x, y, max_width, font_name='Helvetica', font_size=12):
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet

        style = getSampleStyleSheet()['Normal']
        style.fontName = font_name
        style.fontSize = font_size
        story = [Paragraph(text, style)]

        buffer = io.BytesIO()
        # Create a temporary SimpleDocTemplate to calculate text flow
        temp_doc = SimpleDocTemplate(buffer, pagesize=letter,
                                     leftMargin=x, rightMargin=width - x - max_width,
                                     topMargin=height - y, bottomMargin=margin)
        try:
            # Build the story to calculate height
            temp_doc.build(story)
            # We can't easily get line by line position from here, so we'll do basic word wrap
            lines = []
            current_line = ""
            for word in text.split():
                # Estimate width (this is a simplification, reportlab's flowables are better)
                if c.stringWidth(current_line + word + " ", font_name, font_size) < max_width:
                    current_line += word + " "
                else:
                    lines.append(current_line.strip())
                    current_line = word + " "
            lines.append(current_line.strip())

            current_y = y
            for line in lines:
                if current_y < margin:
                    c.showPage()
                    current_y = height - margin
                c.drawString(x, current_y, line)
                current_y -= line_height
            return current_y # Return the final y position
        except Exception as e:
            # Fallback if Paragraph fails or for simplicity
            c.drawString(x, y, "Error rendering content: " + str(e))
            return y - line_height

    c.drawString(margin, y_position, "PBIX Extracted Data")
    y_position -= section_spacing

    for section, content in data.items():
        if y_position < margin:
            c.showPage()
            y_position = height - margin
            c.drawString(margin, y_position, "PBIX Extracted Data (cont.)")
            y_position -= section_spacing

        c.drawString(margin, y_position, f"Section: {section}")
        y_position -= line_height

        content_str = str(content) # Convert all content to string for PDF display

        # Implement simple word wrapping
        lines = []
        current_line = ""
        max_width = width - 2 * margin # Calculate available width
        for word in content_str.split():
             # Add a space before the word unless it's the first word
             test_line = current_line + (" " if current_line else "") + word
             if c.stringWidth(test_line, 'Helvetica', 12) < max_width:
                 current_line = test_line
             else:
                 lines.append(current_line)
                 current_line = word
        lines.append(current_line) # Add the last line

        for line in lines:
            if y_position < margin:
                c.showPage()
                y_position = height - margin
                c.drawString(margin, y_position, f"Section: {section} (cont.)")
                y_position -= line_height
            c.drawString(margin + 20, y_position, line) # Indent content
            y_position -= line_height

        y_position -= section_spacing # Space after section content


    c.save()
    output.seek(0)
    return output


def main():
    st.set_page_config(page_title="PBIX Ray Explorer", layout="wide")

    st.title("PBIX Ray Explorer")

    uploaded_file = handle_file_upload()

    if uploaded_file:
        extracted_data = process_pbix_file(uploaded_file)

        if extracted_data:
            # Display schema details from the 'Schema' attribute provided by PbixUnpacker
            st.header("Schema Details")
            schema_data = extracted_data.get("Schema")

            if schema_data and isinstance(schema_data, dict):
                st.json(schema_data)

                # Attempt to get table names from the Schema data for the dropdown
                schema_names = []
                # The structure of 'schema_data' needs to be inspected.
                # Based on typical schema representations, it might contain a list or dict of tables.
                # Let's assume it's a dictionary where keys are table names or similar identifiers.
                # If it's a list of table objects, we'd need to iterate and access a 'name' attribute.
                # For now, assume it's a dict and keys are names.
                try:
                    schema_names = list(schema_data.keys())
                except Exception as e:
                    st.warning(f"Could not extract schema names from Schema data: {e}")
                    schema_names = [] # Ensure it's an empty list on error


                if schema_names:
                    st.subheader("Select a Schema to View Data")
                    # Data viewing is not supported by pbixray.data_model, so dropdown is for context
                    selected_schema = st.selectbox("Choose a schema", ["--Select--"] + schema_names)

                    if selected_schema and selected_schema != "--Select--":
                         st.subheader(f"Selected Schema Details: {selected_schema}")
                         # Display details for the selected schema if available in schema_data
                         if selected_schema in schema_data:
                             st.json(schema_data[selected_schema])
                         else:
                             st.info(f"Details for schema '{selected_schema}' not found in extracted Schema data.")

                         # Placeholder for data viewing (not supported)
                         st.info(f"Direct data extraction for schema '{selected_schema}' is not supported by the current pbixray library for viewing.")

                else:
                     st.info("No schema names found to display in dropdown.")
            elif schema_data == "Not available":
                 st.info("Schema information is not available in the PBIX file.")
            else:
                st.info("Could not extract Schema details in expected format.")


            # Display other extracted data (excluding Schema, which is handled above)
            display_extracted_data({k: v for k, v in extracted_data.items() if k != "Schema"})

            # Add download buttons
            st.sidebar.header("Download Output")
            if extracted_data:
                # Pass the entire extracted_data to the generators, they can decide what to include
                excel_output = generate_excel_doc(extracted_data)
                st.sidebar.download_button(label="Download Data as Excel",
                                        data=excel_output,
                                        file_name="pbix_data.xlsx",
                                        mime="application/vnd.openxmlformats-officedocument.spreadsheet.sheet")

                word_output = generate_word_doc(extracted_data)
                st.sidebar.download_button(label="Download Data as Word",
                                        data=word_output,
                                        file_name="pbix_data.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                pdf_output = generate_pdf_doc(extracted_data)
                st.sidebar.download_button(label="Download Data as PDF",
                                        data=pdf_output,
                                        file_name="pbix_data.pdf",
                                        mime="application/pdf")


if __name__ == "__main__":
    main()
