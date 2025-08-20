import streamlit as st
import tempfile
import os
import io
import re # Import regex for DAX expression analysis
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import pbixray
import json

def sizeof_fmt(num, suffix='B'):
    """Formats a number into a human-readable byte string."""
    for unit in ['', 'Ki', 'Mi', 'Gi', 'Ti', 'Pi', 'Ei', 'Zi']:
        if abs(num) < 1024.0:
            return f'{num:3.1f}{unit}{suffix}'
        num /= 1024.0
    return f'{num:.1f}Yi{suffix}'

def analyze_pbix_file(pbix_file_path):
    """
    Analyzes a PBIX file and extracts key information.

    Args:
        pbix_file_path: Path to the PBIX file.

    Returns:
        A dictionary containing extracted information about the PBIX file.
    """
    try:
        pbix_model = pbixray.PBIXRay(pbix_file_path)

        analysis_results = {
            "metadata": pbix_model.metadata(),
            "size_bytes": pbix_model.size(),
            "tables": {},
            "relationships": [],
            "measures": {},
            "power_query": pbix_model.power_query(),
            "m_parameters": pbix_model.m_parameters(),
            "dax_tables": pbix_model.dax_tables(),
        }

        # Extract table and column details
        for table in pbix_model.tables():
            table_info = {
                "name": table.name(),
                "row_count": table.row_count(),
                "columns": {}
            }
            for column in table.columns():
                column_info = {
                    "name": column.name(),
                    "data_type": column.data_type(),
                    "is_hidden": column.is_hidden(),
                    "is_keyed": column.is_keyed(),
                    "is_nullable": column.is_nullable(),
                    "is_unique": column.is_unique(),
                    "source_column": column.source_column(),
                    "summarization_function": column.summarization_function(),
                    "data_category": column.data_category(),
                    "display_folder": column.display_folder(),
                    "format_string": column.format_string(),
                    "expression": column.expression() if hasattr(column, 'expression') else None # For calculated columns
                }
                table_info["columns"][column.name()] = column_info
            analysis_results["tables"][table.name()] = table_info

        # Extract measure details
        for measure in pbix_model.measures():
            measure_info = {
                "name": measure.name(),
                "expression": measure.expression(),
                "display_folder": measure.display_folder(),
                "format_string": measure.format_string(),
                "is_hidden": measure.is_hidden()
            }
            analysis_results["measures"][measure.name()] = measure_info

        # Extract relationship details
        for relationship in pbix_model.relationships():
            relationship_info = {
                "name": relationship.name(),
                "from_table": relationship.from_table(),
                "from_column": relationship.from_column(),
                "to_table": relationship.to_table(),
                "to_column": relationship.to_column(),
                "model": relationship.model(), # Indicates relationship type (e.g., 1:M, M:1)
                "active": relationship.is_active(),
                "cross_filter_direction": relationship.cross_filter_direction(),
                "security_filter_table": relationship.security_filter_table()
            }
            analysis_results["relationships"].append(relationship_info)

        return analysis_results

    except Exception as e:
        st.error(f"Error analyzing PBIX file: {e}")
        return None

def generate_documentation(analysis_results, output_format):
    """
    Generates documentation in Word or PDF format from analysis results.

    Args:
        analysis_results: A dictionary containing the analysis results of a PBIX file.
        output_format: The desired output format ('Word' or 'PDF').

    Returns:
        A BytesIO buffer containing the generated document, or None if an error occurs.
    """
    if output_format == 'Word':
        doc = Document()
        doc.add_heading('PBIX File Documentation', 0)

        # Add Metadata
        doc.add_heading('Metadata', level=1)
        metadata = analysis_results.get("metadata", {})
        if metadata:
            for key, value in metadata.items():
                doc.add_paragraph(f"**{key}:** {value}")
        else:
            doc.add_paragraph("No metadata available.")

        # Add Tables
        doc.add_heading('Tables', level=1)
        tables = analysis_results.get("tables", {})
        if tables:
            doc.add_paragraph(f"Number of Tables: {len(tables)}")
            for table_name, table_info in tables.items():
                doc.add_heading(f"Table: {table_name} (Rows: {table_info.get('row_count', 'N/A')})", level=2)
                doc.add_heading('Columns', level=3)
                columns = table_info.get("columns", {})
                if columns:
                    # Create a table for columns
                    column_data = [["Name", "Data Type", "Is Hidden", "Source Column", "Summarization", "Expression"]]
                    for col_name, col_info in columns.items():
                         column_data.append([
                             col_name,
                             col_info.get('data_type', 'N/A'),
                             str(col_info.get('is_hidden', 'N/A')),
                             col_info.get('source_column', 'N/A'),
                             col_info.get('summarization_function', 'N/A'),
                             col_info.get('expression', 'N/A') if col_info.get('expression') else '' # Handle None expression
                         ])
                    if len(column_data) > 1: # Check if there are columns to add to the table
                        table = doc.add_table(rows=len(column_data), cols=len(column_data[0]))
                        table.style = 'Table Grid' # Apply a style for borders
                        for r_idx, row in enumerate(column_data):
                            for c_idx, cell_data in enumerate(row):
                                table.cell(r_idx, c_idx).text = str(cell_data) # Ensure cell data is string
                else:
                    doc.add_paragraph("No columns found for this table.")
        else:
            doc.add_paragraph("No tables found.")

        # Add Measures
        doc.add_heading('Measures', level=1)
        measures = analysis_results.get("measures", {})
        if measures:
            doc.add_paragraph(f"Number of Measures: {len(measures)}")
            for measure_name, measure_info in measures.items():
                doc.add_heading(f"Measure: {measure_name}", level=2)
                doc.add_paragraph(f"**Expression:** {measure_info.get('expression', 'N/A')}")
                doc.add_paragraph(f"**Display Folder:** {measure_info.get('display_folder', 'N/A')}")
                doc.add_paragraph(f"**Format String:** {measure_info.get('format_string', 'N/A')}")
                doc.add_paragraph(f"**Is Hidden:** {measure_info.get('is_hidden', 'N/A')}")
        else:
            doc.add_paragraph("No measures found.")

        # Add Relationships
        doc.add_heading('Relationships', level=1)
        relationships = analysis_results.get("relationships", [])
        if relationships:
            doc.add_paragraph(f"Number of Relationships: {len(relationships)}")
            # Create a table for relationships
            relationship_data = [["Name", "From Table", "From Column", "To Table", "To Column", "Model", "Active", "Cross Filter Direction"]]
            for rel in relationships:
                relationship_data.append([
                    rel.get('name', 'N/A'),
                    rel.get('from_table', 'N/A'),
                    rel.get('from_column', 'N/A'),
                    rel.get('to_table', 'N/A'),
                    rel.get('to_column', 'N/A'),
                    rel.get('model', 'N/A'),
                    str(rel.get('active', 'N/A')),
                    rel.get('cross_filter_direction', 'N/A')
                ])
            if len(relationship_data) > 1: # Check if there are relationships to add to the table
                table = doc.add_table(rows=len(relationship_data), cols=len(relationship_data[0]))
                table.style = 'Table Grid' # Apply a style for borders
                for r_idx, row in enumerate(relationship_data):
                    for c_idx, cell_data in enumerate(row):
                         table.cell(r_idx, c_idx).text = str(cell_data) # Ensure cell data is string
        else:
            doc.add_paragraph("No relationships found.")

        # Add Power Query (M Code)
        doc.add_heading('Power Query (M Code)', level=1)
        power_query = analysis_results.get("power_query", "No M code found.")
        if power_query and power_query.strip():
            doc.add_paragraph("M Code:")
            doc.add_paragraph(power_query) # Add M code as a block of text
        else:
            doc.add_paragraph("No Power Query (M Code) found.")

        # Add M Parameters
        doc.add_heading('M Parameters', level=1)
        m_parameters = analysis_results.get("m_parameters", [])
        if m_parameters:
             doc.add_paragraph(f"Number of M Parameters: {len(m_parameters)}")
             # Add parameters as a list or formatted text
             for param in m_parameters:
                 doc.add_paragraph(f"- Name: {param.get('name', 'N/A')}, Value: {param.get('value', 'N/A')}, Data Type: {param.get('dataType', 'N/A')}")
        else:
            doc.add_paragraph("No M Parameters found.")

        # Add DAX Tables
        doc.add_heading('DAX Tables', level=1)
        dax_tables = analysis_results.get("dax_tables", {})
        if dax_tables:
            doc.add_paragraph(f"Number of DAX Tables: {len(dax_tables)}")
            for dax_table_name, dax_table_info in dax_tables.items():
                 doc.add_heading(f"DAX Table: {dax_table_name}", level=2)
                 doc.add_paragraph(f"**Expression:** {dax_table_info.get('expression', 'N/A')}")
                 doc.add_paragraph(f"**Display Folder:** {dax_table_info.get('display_folder', 'N/A')}")
                 doc.add_paragraph(f"**Is Hidden:** {dax_table_info.get('is_hidden', 'N/A')}")
        else:
            doc.add_paragraph("No DAX Tables found.")


        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    elif output_format == 'PDF':
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Add Title
        story.append(Paragraph("PBIX File Documentation", styles['h1']))
        story.append(Spacer(1, 12))

        # Add Metadata
        story.append(Paragraph("Metadata", styles['h2']))
        metadata = analysis_results.get("metadata", {})
        if metadata:
            for key, value in metadata.items():
                story.append(Paragraph(f"<b>{key}:</b> {value}", styles['Normal']))
        else:
            story.append(Paragraph("No metadata available.", styles['Normal']))
        story.append(Spacer(1, 12))

        # Add Tables
        story.append(Paragraph("Tables", styles['h2']))
        tables = analysis_results.get("tables", {})
        if tables:
            story.append(Paragraph(f"Number of Tables: {len(tables)}", styles['Normal']))
            for table_name, table_info in tables.items():
                story.append(Paragraph(f"Table: {table_name} (Rows: {table_info.get('row_count', 'N/A')})", styles['h3']))
                story.append(Paragraph("Columns", styles['h4']))
                columns = table_info.get("columns", {})
                if columns:
                    column_data = [["Name", "Data Type", "Is Hidden", "Source Column", "Summarization", "Expression"]]
                    for col_name, col_info in columns.items():
                         column_data.append([
                             col_name,
                             col_info.get('data_type', 'N/A'),
                             str(col_info.get('is_hidden', 'N/A')),
                             col_info.get('source_column', 'N/A'),
                             col_info.get('summarization_function', 'N/A'),
                             col_info.get('expression', 'N/A') if col_info.get('expression') else ''
                         ])
                    if len(column_data) > 1:
                        table = Table(column_data)
                        table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, 0), 10),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ]))
                        story.append(table)
                else:
                    story.append(Paragraph("No columns found for this table.", styles['Normal']))
                story.append(Spacer(1, 6)) # Small spacer between tables
        else:
            story.append(Paragraph("No tables found.", styles['Normal']))
        story.append(Spacer(1, 12))

        # Add Measures
        story.append(Paragraph("Measures", styles['h2']))
        measures = analysis_results.get("measures", {})
        if measures:
            story.append(Paragraph(f"Number of Measures: {len(measures)}", styles['Normal']))
            for measure_name, measure_info in measures.items():
                story.append(Paragraph(f"Measure: {measure_name}", styles['h3']))
                story.append(Paragraph(f"<b>Expression:</b> {measure_info.get('expression', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Display Folder:</b> {measure_info.get('display_folder', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Format String:</b> {measure_info.get('format_string', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Is Hidden:</b> {measure_info.get('is_hidden', 'N/A')}", styles['Normal']))
                story.append(Spacer(1, 6))
        else:
            story.append(Paragraph("No measures found.", styles['Normal']))
        story.append(Spacer(1, 12))

        # Add Relationships
        story.append(Paragraph("Relationships", styles['h2']))
        relationships = analysis_results.get("relationships", [])
        if relationships:
            story.append(Paragraph(f"Number of Relationships: {len(relationships)}", styles['Normal']))
            relationship_data = [["Name", "From Table", "From Column", "To Table", "To Column", "Model", "Active", "Cross Filter Direction"]]
            for rel in relationships:
                relationship_data.append([
                    rel.get('name', 'N/A'),
                    rel.get('from_table', 'N/A'),
                    rel.get('from_column', 'N/A'),
                    rel.get('to_table', 'N/A'),
                    rel.get('to_column', 'N/A'),
                    rel.get('model', 'N/A'),
                    str(rel.get('active', 'N/A')),
                    rel.get('cross_filter_direction', 'N/A')
                ])
            if len(relationship_data) > 1:
                 table = Table(relationship_data)
                 table.setStyle(TableStyle([
                     ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                     ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                     ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                     ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                     ('FONTSIZE', (0, 0), (-1, 0), 10),
                     ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                     ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                     ('GRID', (0, 0), (-1, -1), 1, colors.black),
                 ]))
                 story.append(table)
        else:
            story.append(Paragraph("No relationships found.", styles['Normal']))
        story.append(Spacer(1, 12))

        # Add Power Query (M Code)
        story.append(Paragraph("Power Query (M Code)", styles['h2']))
        power_query = analysis_results.get("power_query", "No M code found.")
        if power_query and power_query.strip():
            story.append(Paragraph("M Code:", styles['Normal']))
            story.append(Paragraph(power_query, styles['Code'])) # Use Code style for M code
        else:
            story.append(Paragraph("No Power Query (M Code) found.", styles['Normal']))
        story.append(Spacer(1, 12))

        # Add M Parameters
        story.append(Paragraph("M Parameters", styles['h2']))
        m_parameters = analysis_results.get("m_parameters", [])
        if m_parameters:
            story.append(Paragraph(f"Number of M Parameters: {len(m_parameters)}", styles['Normal']))
            for param in m_parameters:
                story.append(Paragraph(f"- Name: {param.get('name', 'N/A')}, Value: {param.get('value', 'N/A')}, Data Type: {param.get('dataType', 'N/A')}", styles['Normal']))
        else:
            story.append(Paragraph("No M Parameters found.", styles['Normal']))
        story.append(Spacer(1, 12))

        # Add DAX Tables
        story.append(Paragraph("DAX Tables", styles['h2']))
        dax_tables = analysis_results.get("dax_tables", {})
        if dax_tables:
            story.append(Paragraph(f"Number of DAX Tables: {len(dax_tables)}", styles['Normal']))
            for dax_table_name, dax_table_info in dax_tables.items():
                 story.append(Paragraph(f"DAX Table: {dax_table_name}", styles['h3']))
                 story.append(Paragraph(f"<b>Expression:</b> {dax_table_info.get('expression', 'N/A')}", styles['Normal']))
                 story.append(Paragraph(f"<b>Display Folder:</b> {dax_table_info.get('display_folder', 'N/A')}", styles['Normal']))
                 story.append(Paragraph(f"<b>Is Hidden:</b> {dax_table_info.get('is_hidden', 'N/A')}", styles['Normal']))
                 story.append(Spacer(1, 6))
        else:
            story.append(Paragraph("No DAX Tables found.", styles['Normal']))
        story.append(Spacer(1, 12))


        try:
            doc.build(story)
            buffer.seek(0)
            return buffer
        except Exception as e:
            st.error(f"Error building PDF document: {e}")
            return None


    else:
        st.error(f"Unsupported output format: {output_format}")
        return None

def analyze_performance_and_complexity(analysis_results):
    """
    Analyzes the PBIX analysis results for potential performance issues and complexity.

    Args:
        analysis_results: A dictionary containing the analysis results of a PBIX file.

    Returns:
        A dictionary containing findings related to performance and complexity.
    """
    performance_insights = {
        "complex_measures": [],
        "complex_calculated_columns": [],
        "many_to_many_without_bridge": [],
        "large_tables": [],
        "tables_with_many_columns": []
    }

    # Thresholds (can be adjusted)
    COMPLEX_DAX_THRESHOLD = 100 # Arbitrary length/complexity indicator
    LARGE_TABLE_ROW_THRESHOLD = 100000 # Example: Tables with over 100k rows
    MANY_COLUMNS_THRESHOLD = 50 # Example: Tables with over 50 columns

    # Analyze Measures for complexity
    measures = analysis_results.get("measures", {})
    for measure_name, measure_info in measures.items():
        expression = measure_info.get("expression", "")
        # Simple check for expression length as an indicator of complexity
        if len(expression) > COMPLEX_DAX_THRESHOLD:
            performance_insights["complex_measures"].append({
                "name": measure_name,
                "expression_length": len(expression)
            })
        # Add more sophisticated DAX analysis here if possible (e.g., checking for specific functions)

    # Analyze Calculated Columns for complexity and potential impact
    tables = analysis_results.get("tables", {})
    for table_name, table_info in tables.items():
        columns = table_info.get("columns", {})
        for col_name, col_info in columns.items():
            if col_info.get("expression"): # Check if it's a calculated column
                expression = col_info.get("expression")
                # Simple check for expression length
                if len(expression) > COMPLEX_DAX_THRESHOLD:
                    performance_insights["complex_calculated_columns"].append({
                        "table": table_name,
                        "column": col_name,
                        "expression_length": len(expression)
                    })
                # Add more sophisticated DAX analysis for calculated columns

        # Identify large tables
        row_count = table_info.get("row_count", 0)
        if isinstance(row_count, (int, float)) and row_count > LARGE_TABLE_ROW_THRESHOLD:
             performance_insights["large_tables"].append({
                 "name": table_name,
                 "row_count": row_count
             })

        # Identify tables with many columns
        if columns and len(columns) > MANY_COLUMNS_THRESHOLD:
             performance_insights["tables_with_many_columns"].append({
                 "name": table_name,
                 "column_count": len(columns)
             })


    # Analyze Relationships for potential many-to-many issues without bridge tables
    # This is a simplified check. A proper check would require analyzing cardinality
    # and checking for dedicated bridge tables.
    relationships = analysis_results.get("relationships", [])
    # Simplified check: Look for relationships explicitly marked as Many-to-Many by pbixray
    # or infer based on model type (though pbixray's 'model' might not be explicit M:M)
    # A more robust check would involve analyzing table cardinality and relationship direction.
    # For this example, we'll just note relationships that *might* be M:M based on simplified criteria
    # and highlight the need for bridge tables if they are not present.
    # The 'model' attribute from pbixray doesn't directly give M:M.
    # A true M:M detection requires checking cardinality on both sides, which pbixray's basic model
    # might not expose directly in a simple way.
    # Let's skip the many-to-many detection for now as it requires deeper model analysis
    # not readily available from the current analysis_results structure.

    return performance_insights


def st_app():
    """Main Streamlit application function for PBIX Analyzer."""
    st.set_page_config(page_title="PBIX Analyzer", layout="wide")
    st.title("PBIX Analyzer")

    uploaded_file = st.file_uploader("Upload a PBIX file", type=["pbix"])

    if uploaded_file is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pbix") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name

        try:
            st.info("Analyzing PBIX file...")
            analysis_results = analyze_pbix_file(tmp_file_path)

            if analysis_results:
                st.success("Analysis complete!")

                # Display analysis results
                st.header("Metadata")
                st.write(f"**File Size:** {sizeof_fmt(analysis_results.get('size_bytes', 0))}")
                if analysis_results.get("metadata"):
                    st.json(analysis_results["metadata"])

                st.header("Tables")
                tables = analysis_results.get("tables", {})
                if tables:
                    st.write(f"Number of Tables: {len(tables)}")
                    for table_name, table_info in tables.items():
                        with st.expander(f"Table: {table_name} (Rows: {table_info.get('row_count', 'N/A')})"):
                            st.subheader("Columns")
                            columns = table_info.get("columns", {})
                            if columns:
                                column_data = [{"Name": col_name, **col_info} for col_name, col_info in columns.items()]
                                st.dataframe(column_data)
                            else:
                                st.write("No columns found for this table.")

                st.header("Measures")
                measures = analysis_results.get("measures", {})
                if measures:
                    st.write(f"Number of Measures: {len(measures)}")
                    for measure_name, measure_info in measures.items():
                         with st.expander(f"Measure: {measure_name}"):
                            st.write(f"**Expression:**")
                            st.code(measure_info.get('expression', 'N/A'), language='dax')
                            st.write(f"**Display Folder:** {measure_info.get('display_folder', 'N/A')}")
                            st.write(f"**Format String:** {measure_info.get('format_string', 'N/A')}")
                            st.write(f"**Is Hidden:** {measure_info.get('is_hidden', 'N/A')}")


                st.header("Relationships")
                relationships = analysis_results.get("relationships", [])
                if relationships:
                    st.write(f"Number of Relationships: {len(relationships)}")
                    relationship_data = []
                    for rel in relationships:
                        relationship_data.append({
                            "Name": rel.get('name', 'N/A'),
                            "From Table": rel.get('from_table', 'N/A'),
                            "From Column": rel.get('from_column', 'N/A'),
                            "To Table": rel.get('to_table', 'N/A'),
                            "To Column": rel.get('to_column', 'N/A'),
                            "Model": rel.get('model', 'N/A'),
                            "Active": rel.get('active', 'N/A'),
                            "Cross Filter Direction": rel.get('cross_filter_direction', 'N/A'),
                            "Security Filter Table": rel.get('security_filter_table', 'N/A')
                        })
                    st.dataframe(relationship_data)
                else:
                    st.write("No relationships found.")

                st.header("Power Query (M Code)")
                power_query = analysis_results.get("power_query", "No M code found.")
                if power_query and power_query.strip():
                     with st.expander("View M Code"):
                        st.code(power_query, language='m')
                else:
                    st.write("No Power Query (M Code) found.")


                st.header("M Parameters")
                m_parameters = analysis_results.get("m_parameters", [])
                if m_parameters:
                    st.write(f"Number of M Parameters: {len(m_parameters)}")
                    st.json(m_parameters)
                else:
                    st.write("No M Parameters found.")

                st.header("DAX Tables")
                dax_tables = analysis_results.get("dax_tables", {})
                if dax_tables:
                    st.write(f"Number of DAX Tables: {len(dax_tables)}")
                    for dax_table_name, dax_table_info in dax_tables.items():
                        with st.expander(f"DAX Table: {dax_table_name}"):
                             st.write(f"**Expression:**")
                             st.code(dax_table_info.get('expression', 'N/A'), language='dax')
                             st.write(f"**Display Folder:** {dax_table_info.get('display_folder', 'N/A')}")
                             st.write(f"**Is Hidden:** {dax_table_info.get('is_hidden', 'N/A')}")
                else:
                    st.write("No DAX Tables found.")


                # Performance and Complexity Analysis Section
                st.header("Performance and Complexity Insights")
                performance_insights = analyze_performance_and_complexity(analysis_results)

                if any(performance_insights.values()): # Check if there are any findings
                    st.subheader("Potential Performance Issues and Complexity:")

                    if performance_insights["complex_measures"]:
                        st.warning("Complex Measures Identified:")
                        for measure in performance_insights["complex_measures"]:
                            st.write(f"- Measure: **{measure['name']}** (Expression Length: {measure['expression_length']})")

                    if performance_insights["complex_calculated_columns"]:
                        st.warning("Complex Calculated Columns Identified:")
                        for col in performance_insights["complex_calculated_columns"]:
                            st.write(f"- Table: **{col['table']}**, Column: **{col['column']}** (Expression Length: {col['expression_length']})")

                    if performance_insights["large_tables"]:
                         st.info("Large Tables Identified:")
                         for table in performance_insights["large_tables"]:
                             st.write(f"- Table: **{table['name']}** (Row Count: {table['row_count']})")

                    if performance_insights["tables_with_many_columns"]:
                         st.info("Tables with Many Columns Identified:")
                         for table in performance_insights["tables_with_many_columns"]:
                             st.write(f"- Table: **{table['name']}** (Column Count: {table['column_count']})")

                else:
                    st.info("No significant performance or complexity issues identified based on current checks.")


                # Documentation Generation Section
                st.header("Generate Documentation")
                output_format = st.radio("Select Output Format", ('Word', 'PDF'))

                if st.button("Generate Documentation"):
                    st.info(f"Generating {output_format} documentation...")
                    doc_buffer = generate_documentation(analysis_results, output_format)

                    if doc_buffer:
                        file_name = f"pbix_documentation.{'docx' if output_format == 'Word' else 'pdf'}"
                        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if output_format == 'Word' else 'application/pdf'

                        st.download_button(
                            label=f"Download {output_format} Document",
                            data=doc_buffer,
                            file_name=file_name,
                            mime=mime_type
                        )
                        st.success(f"{output_format} documentation generated successfully!")
                    else:
                        st.error(f"Failed to generate {output_format} documentation.")


            else:
                st.error("Failed to analyze the PBIX file. Please ensure it is a valid PBIX file.")

        except Exception as e:
            st.error(f"An error occurred during analysis or documentation generation: {e}")
        finally:
            # Clean up the temporary file
            if os.path.exists(tmp_file_path):
                os.remove(tmp_file_path)

    else:
        st.info("Please upload a PBIX file to analyze.")

if __name__ == "__main__":
    st_app() # Call the main Streamlit application function
